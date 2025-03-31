import io
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render
from django.urls import reverse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE
from .models import Keyword
from .forms import DocumentUploadForm
from django.shortcuts import render, get_object_or_404
from doc_processor.models import Keyword


#Home page
def home(request):
    return render(request, 'home.html')

def add_hyperlink(paragraph, keyword, url):
    """
    Adds a hyperlink to the keyword in the paragraph and changes its color.
    """
    # Define a new run for hyperlink text
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Iterate through existing runs in the paragraph
    for run in paragraph.runs:
        if keyword in run.text:
            # Split text around the keyword
            parts = run.text.split(keyword)
            run.text = parts[0]  # Text before the keyword
            
            # Add runs before and after the hyperlink
            if parts[0]:
                paragraph.add_run(parts[0])  # Text before the hyperlink

            # Create a run for the hyperlink
            hyperlink_run = paragraph.add_run(keyword)
            hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
            
            # Add the hyperlink to the XML
            hyperlink_run_element = hyperlink_run._element
            hyperlink.append(hyperlink_run_element)
            
            # Append the hyperlink to the paragraph
            paragraph._element.append(hyperlink)
            
            # Add the remaining parts of the text
            if len(parts) > 1:
                paragraph.add_run(parts[1])  # Text after the hyperlink
            break  # Only process the keyword once per run
def process_document(request):
    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['document']
            doc = Document(uploaded_file)

            # Fetch all keywords from the database
            keywords = Keyword.objects.all()

            # Process the document to add hyperlinks and change color
            for para in doc.paragraphs:
                for keyword in keywords:
                    if keyword.word in para.text:
                        # Create the full URL to be attached
                        full_url = f"http://localhost:8000/doc_processor/description/{keyword.id}/"
                        add_hyperlink(para, keyword.word, full_url)

            # Save the modified document to a BytesIO object
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            # Return the response with the modified document
            response = HttpResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{uploaded_file.name}"'
            return response
    else:
        form = DocumentUploadForm()

    return render(request, 'doc_processor/upload.html', {'form': form})


def keyword_description(request, keyword_id):
    keyword = get_object_or_404(Keyword, id=keyword_id)
    return render(request, 'doc_processor/keyword_description.html', {'keyword': keyword})
